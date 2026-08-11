import os
import glob
import hashlib
from typing import List, Dict, Any
from pathlib import Path
from datetime import datetime

from qdrant_client import QdrantClient
from qdrant_client.models import PointStruct
from langchain_ollama import OllamaEmbeddings

# ✅ CORREGIDO: Importar desde langchain_text_splitters
try:
    from langchain_text_splitters import RecursiveCharacterTextSplitter
except ImportError:
    # Fallback para versiones antiguas
    from langchain.text_splitter import RecursiveCharacterTextSplitter

from app.config import settings
from app.rag.vector_store import ensure_vector_collection


class DocumentIngestor:
    """
    Ingestor de documentos técnicos para la base de conocimiento del agente.
    Soporta: PDF, TXT, DOCX, Markdown, JSON, CSV
    """
    
    def __init__(self):
        self.embeddings = OllamaEmbeddings(
            base_url=settings.OLLAMA_BASE_URL,
            model=settings.EMBEDDING_MODEL_NAME
        )
        self.qdrant = QdrantClient(url=settings.VECTOR_DB_URL)
        self.collection = settings.VECTOR_COLLECTION_NAME
        
        # Verificar que la colección existe, si no crearla
        self._ensure_collection()
        
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=1000,
            chunk_overlap=200,
            separators=["\n\n", "\n", ". ", " ", ""]
        )
    
    def _ensure_collection(self):
        """Asegura que la colección existe en Qdrant."""
        try:
            ensure_vector_collection(self.qdrant, self.collection, self.embeddings)
        except Exception as e:
            print(f"⚠️ Error verificando colección: {e}")
    
    def ingest_file(self, file_path: str, metadata: Dict[str, Any] = None) -> int:
        """
        Ingiere un archivo completo en la base vectorial.
        
        Args:
            file_path: Ruta al archivo
            metadata: Metadatos adicionales (categoría, fuente, etc.)
        
        Returns:
            Número de fragmentos indexados
        """
        if not os.path.exists(file_path):
            print(f"❌ Archivo no encontrado: {file_path}")
            return 0
        
        # Detectar tipo de archivo por extensión
        ext = Path(file_path).suffix.lower()
        
        print(f"📄 Procesando: {os.path.basename(file_path)} ({ext})")
        
        if ext == '.txt':
            content = self._read_txt(file_path)
        elif ext == '.pdf':
            content = self._read_pdf(file_path)
        elif ext == '.docx':
            content = self._read_docx(file_path)
        elif ext == '.md':
            content = self._read_markdown(file_path)
        elif ext == '.json':
            content = self._read_json(file_path)
        elif ext == '.csv':
            content = self._read_csv(file_path)
        else:
            print(f"⚠️ Extensión no soportada: {ext}")
            return 0
        
        if not content or len(content.strip()) < 50:
            print(f"⚠️ El archivo está vacío o es muy corto: {file_path}")
            return 0
        
        # Dividir en fragmentos
        chunks = self.text_splitter.split_text(content)
        print(f"   📝 Dividido en {len(chunks)} fragmentos")
        
        # Metadatos base
        base_metadata = {
            "source": os.path.basename(file_path),
            "file_path": file_path,
            "file_type": ext[1:] if ext else "unknown",
            "chunk_count": len(chunks),
            "timestamp": datetime.now().isoformat()
        }
        if metadata:
            base_metadata.update(metadata)
        
        # Indexar cada fragmento
        points = []
        for i, chunk in enumerate(chunks):
            if len(chunk.strip()) < 50:  # Saltar fragmentos muy cortos
                continue
            
            # Generar embedding
            try:
                vector = self.embeddings.embed_query(chunk)
            except Exception as e:
                print(f"   ⚠️ Error generando embedding: {e}")
                continue
            
            # ID basado en hash del contenido
            content_hash = hashlib.md5(chunk.encode()).hexdigest()[:16]
            point_id = f"doc_{content_hash}_{i}"
            
            point = PointStruct(
                id=point_id,
                vector=vector,
                payload={
                    "page_content": chunk,
                    "chunk_index": i,
                    **base_metadata
                }
            )
            points.append(point)
        
        if points:
            try:
                self.qdrant.upsert(
                    collection_name=self.collection,
                    points=points
                )
                print(f"   ✅ Indexados {len(points)} fragmentos")
                return len(points)
            except Exception as e:
                print(f"   ❌ Error indexando: {e}")
                return 0
        
        return 0
    
    def ingest_directory(self, directory: str, pattern: str = "*", metadata: Dict[str, Any] = None) -> int:
        """
        Ingiere todos los archivos de un directorio.
        
        Args:
            directory: Ruta al directorio
            pattern: Patrón de archivos (ej: "*.pdf", "*.txt")
            metadata: Metadatos para todos los archivos
        """
        if not os.path.exists(directory):
            print(f"❌ Directorio no encontrado: {directory}")
            return 0
        
        files = glob.glob(os.path.join(directory, pattern))
        if not files:
            print(f"⚠️ No se encontraron archivos en {directory} con patrón {pattern}")
            return 0
        
        print(f"📂 Encontrados {len(files)} archivos en {directory}")
        
        total = 0
        for file_path in files:
            total += self.ingest_file(file_path, metadata)
        
        print(f"   Total indexado: {total} fragmentos")
        return total
    
    # ============================================================
    # LECTURA DE DIFERENTES FORMATOS
    # ============================================================
    
    def _read_txt(self, file_path: str) -> str:
        """Lee un archivo de texto plano."""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                return f.read()
        except UnicodeDecodeError:
            with open(file_path, 'r', encoding='latin-1') as f:
                return f.read()
        except Exception as e:
            print(f"   ⚠️ Error leyendo TXT: {e}")
            return ""
    
    def _read_pdf(self, file_path: str) -> str:
        """Lee un archivo PDF."""
        try:
            import PyPDF2
            content = []
            with open(file_path, 'rb') as f:
                reader = PyPDF2.PdfReader(f)
                for page in reader.pages:
                    text = page.extract_text()
                    if text:
                        content.append(text)
            return "\n\n".join(content)
        except ImportError:
            print("   ⚠️ PyPDF2 no está instalado. Instala: pip install PyPDF2")
            return ""
        except Exception as e:
            print(f"   ⚠️ Error leyendo PDF: {e}")
            return ""
    
    def _read_docx(self, file_path: str) -> str:
        """Lee un archivo DOCX."""
        try:
            import docx
            doc = docx.Document(file_path)
            content = []
            for para in doc.paragraphs:
                if para.text.strip():
                    content.append(para.text)
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        content.append(" | ".join(row_text))
            return "\n\n".join(content)
        except ImportError:
            print("   ⚠️ python-docx no está instalado. Instala: pip install python-docx")
            return ""
        except Exception as e:
            print(f"   ⚠️ Error leyendo DOCX: {e}")
            return ""
    
    def _read_markdown(self, file_path: str) -> str:
        """Lee un archivo Markdown."""
        return self._read_txt(file_path)
    
    def _read_json(self, file_path: str) -> str:
        """Lee un archivo JSON y lo convierte a texto."""
        try:
            import json
            with open(file_path, 'r', encoding='utf-8') as f:
                data = json.load(f)
            
            if isinstance(data, list):
                text = "Lista de elementos:\n"
                for i, item in enumerate(data[:20]):  # Limitar a 20 items
                    text += f"Item {i+1}: {json.dumps(item, ensure_ascii=False, indent=2)}\n"
                return text
            elif isinstance(data, dict):
                return json.dumps(data, ensure_ascii=False, indent=2)
            else:
                return str(data)
        except Exception as e:
            print(f"   ⚠️ Error leyendo JSON: {e}")
            return ""
    
    def _read_csv(self, file_path: str) -> str:
        """Lee un archivo CSV."""
        try:
            import pandas as pd
            df = pd.read_csv(file_path)
            text = f"Archivo CSV: {os.path.basename(file_path)}\n"
            text += f"Columnas: {', '.join(df.columns.tolist())}\n"
            text += f"Filas: {len(df)}\n\n"
            text += df.to_string(max_rows=20, max_cols=10)
            return text
        except ImportError:
            print("   ⚠️ pandas no está instalado. Instala: pip install pandas")
            return ""
        except Exception as e:
            print(f"   ⚠️ Error leyendo CSV: {e}")
            return ""


# ============================================================
# FUNCIÓN DE INGESTA RÁPIDA
# ============================================================

def ingest_knowledge_base(documents_dir: str = "./knowledge_base"):
    """
    Ingiere toda la base de conocimiento desde un directorio.
    """
    if not os.path.exists(documents_dir):
        print(f"❌ Directorio no encontrado: {documents_dir}")
        print("   Crea la estructura: knowledge_base/ con tus documentos")
        return
    
    ingestor = DocumentIngestor()
    
    # Ingesta de todos los archivos recursivamente
    print(f"\n📂 Indexando documentos desde: {documents_dir}")
    total = ingestor.ingest_directory(
        documents_dir,
        pattern="*.*",
        metadata={"source": "knowledge_base", "language": "es"}
    )
    
    print(f"\n✅ Ingesta completada. Total: {total} fragmentos indexados")
    return total


# ============================================================
# SCRIPT DE EJECUCIÓN
# ============================================================

if __name__ == "__main__":
    print("\n" + "="*60)
    print("📚 INGESTOR DE DOCUMENTOS PARA BASE DE CONOCIMIENTO")
    print("="*60 + "\n")
    
    # Cambiar al directorio del proyecto si es necesario
    os.chdir(os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__)))))
    
    # Ejecutar ingesta
    ingest_knowledge_base("./knowledge_base")
