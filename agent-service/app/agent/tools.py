import json
import os
import re
import threading
from time import perf_counter, time
from typing import Any, Optional, Union
import uuid
import hashlib
from datetime import datetime
from pathlib import Path
from urllib.parse import urlparse

import httpx
import pandas as pd
from langchain_core.tools import tool
from langchain_ollama import OllamaEmbeddings
import pymssql
from qdrant_client import QdrantClient
from qdrant_client.models import PointStruct, Distance, VectorParams

from app.config import settings
from app.connectors.db_client import (
    get_solidset_login_for_active_agent,
    list_active_solidset_instances,
)
from app.rag.vector_store import ensure_vector_collection

from app.rag.audio_processor import extract_audio_features
from app.rag.retriever import get_rag_context


def _resolve_solidset_meeting_id(
    meeting_id: Optional[str],
    channel_id: Optional[str],
    meeting_code: Optional[str] = None,
) -> Optional[str]:
    """Devuelve un meeting existente y perteneciente al canal indicado."""
    try:
        channel_uuid = uuid.UUID(str(channel_id or "").strip())
    except (ValueError, TypeError, AttributeError):
        return None
    try:
        requested_uuid = uuid.UUID(str(meeting_id or "").strip())
    except (ValueError, TypeError, AttributeError):
        requested_uuid = None
    code = str(meeting_code or "").strip()
    try:
        with pymssql.connect(
            **settings.sql_server_connection_options(),
            user=settings.SQL_SERVER_USER,
            password=settings.SQL_SERVER_PASSWORD,
            database=settings.SQL_SERVER_DB,
            login_timeout=max(3, settings.DB_INGEST_CONNECT_TIMEOUT_SECONDS),
            timeout=max(10, settings.DB_INGEST_CONNECT_TIMEOUT_SECONDS),
        ) as connection:
            cursor = connection.cursor(as_dict=True)
            if requested_uuid is not None:
                cursor.execute(
                    '''
                    SELECT TOP 1 ID
                    FROM dbo.SysMeeting WITH (NOLOCK)
                    WHERE ID = %s AND IDChannel = %s AND Active = 1
                    ''',
                    (requested_uuid, channel_uuid),
                )
                row = cursor.fetchone()
                if row:
                    return str(row["ID"])
            if code:
                cursor.execute(
                    '''
                    SELECT TOP 1 ID
                    FROM dbo.SysMeeting WITH (NOLOCK)
                    WHERE IDChannel = %s AND Code = %s AND Active = 1
                    ORDER BY ModifiedTime DESC, CreationDate DESC
                    ''',
                    (channel_uuid, code),
                )
                row = cursor.fetchone()
                if row:
                    return str(row["ID"])
    except pymssql.Error as exc:
        print(f"⚠️ No se pudo validar el meeting de SolidSET: {exc}")
    return None


def _generated_docs_dir() -> Path:
    target = Path(settings.GENERATED_DOCS_DIR).expanduser()
    if not target.is_absolute():
        root = Path(__file__).resolve().parents[2]
        target = (root / target).resolve()
    target.mkdir(parents=True, exist_ok=True)
    return target


def _safe_file_stem(value: Optional[str], fallback: str) -> str:
    raw = (value or "").strip()
    if not raw:
        raw = fallback
    clean = re.sub(r"[^A-Za-z0-9_\-\s]", "", raw).strip().replace(" ", "_")
    clean = re.sub(r"_+", "_", clean)
    return clean[:60] or fallback


def _timestamp_suffix() -> str:
    return datetime.now().strftime("%Y%m%d_%H%M%S")


def _solidset_candidate_base_urls(base_url: str) -> list[str]:
    if not base_url:
        return []

    urls = [base_url.rstrip("/")]
    parsed = urlparse(base_url)
    hostname = (parsed.hostname or "").lower()
    running_in_container = os.path.exists("/.dockerenv") or os.getenv("RUNNING_IN_DOCKER") == "1"
    if running_in_container and hostname in {"localhost", "127.0.0.1"}:
        alt = base_url.replace("localhost", "host.docker.internal").replace("127.0.0.1", "host.docker.internal")
        alt = alt.rstrip("/")
        if alt not in urls:
            # Dentro del contenedor localhost nunca es el host Windows. Se
            # prueba primero la traducción para no consumir un timeout inútil.
            urls.insert(0, alt)
    return urls


def _solidset_action_headers(host_header: Optional[str] = None) -> dict:
    headers = {
        "Accept": "application/json",
        "X-Requested-With": "XMLHttpRequest",
    }
    if str(host_header or "").strip():
        headers["Host"] = str(host_header).strip()

    cookie_parts = []
    if settings.SOLIDSET_WORKSTATION_ID:
        cookie_parts.append(f"IDWorkstation={settings.SOLIDSET_WORKSTATION_ID}")
    if settings.SOLIDSET_WORKSTATION_NAME:
        cookie_parts.append(f"NameWorkstation={settings.SOLIDSET_WORKSTATION_NAME}")
    if settings.SOLIDSET_CLIENT_VERSION:
        cookie_parts.append(f"ClientVersion={settings.SOLIDSET_CLIENT_VERSION}")
    if settings.SOLIDSET_APPLICATION_ID:
        cookie_parts.append(f"IDApplication={settings.SOLIDSET_APPLICATION_ID}")
    if cookie_parts:
        headers["Cookie"] = ";".join(cookie_parts)

    return headers


_solidset_runtime_lock = threading.Lock()
_solidset_runtime_auth: dict[str, Any] = {
    "base_url": "",
    "access_key": "",
    "cookie_header": "",
    "authenticated_at": 0.0,
    "login_endpoint": "",
}


def _solidset_cookie_dict_from_header(cookie_header: str) -> dict[str, str]:
    cookie_map: dict[str, str] = {}
    for chunk in (cookie_header or "").split(";"):
        part = chunk.strip()
        if not part or "=" not in part:
            continue
        key, value = part.split("=", 1)
        key = key.strip()
        if not key:
            continue
        cookie_map[key] = value.strip()
    return cookie_map


def _solidset_cookie_header_from_dict(cookie_map: dict[str, str]) -> str:
    if not cookie_map:
        return ""
    return "; ".join(f"{k}={v}" for k, v in cookie_map.items() if k)


def _solidset_merge_cookie_headers(*headers: str) -> str:
    merged: dict[str, str] = {}
    for header in headers:
        merged.update(_solidset_cookie_dict_from_header(header))
    return _solidset_cookie_header_from_dict(merged)


def _solidset_extract_access_key(response: httpx.Response) -> str:
    content_type = (response.headers.get("content-type") or "").lower()
    if "json" not in content_type:
        return ""

    try:
        payload = response.json()
    except Exception:
        return ""

    if not isinstance(payload, dict):
        return ""

    access_key = payload.get("accessKey") or payload.get("AccessKey") or ""
    return access_key.strip() if isinstance(access_key, str) else ""


def _solidset_client_cookie_header(client: httpx.Client) -> str:
    try:
        cookie_map = {name: value for name, value in client.cookies.items()}
    except Exception:
        cookie_map = {}
    return _solidset_cookie_header_from_dict(cookie_map)


def _solidset_set_runtime_auth(
    *,
    base_url: str,
    cookie_header: str,
    access_key: str,
    login_endpoint: str,
) -> None:
    with _solidset_runtime_lock:
        _solidset_runtime_auth["base_url"] = (base_url or "").rstrip("/")
        _solidset_runtime_auth["cookie_header"] = cookie_header or ""
        _solidset_runtime_auth["access_key"] = access_key or ""
        _solidset_runtime_auth["login_endpoint"] = login_endpoint or ""
        _solidset_runtime_auth["authenticated_at"] = time()


def _solidset_get_runtime_auth() -> dict[str, Any]:
    with _solidset_runtime_lock:
        return dict(_solidset_runtime_auth)


def _solidset_clear_runtime_auth() -> None:
    with _solidset_runtime_lock:
        _solidset_runtime_auth.update(
            {
                "base_url": "",
                "access_key": "",
                "cookie_header": "",
                "authenticated_at": 0.0,
                "login_endpoint": "",
            }
        )


def _solidset_get_all_base_candidates() -> list[str]:
    collected: list[str] = []
    try:
        configured_urls = [row.get("BaseUrl") for row in list_active_solidset_instances()]
    except Exception:
        configured_urls = []
    for configured in configured_urls:
        for candidate in _solidset_candidate_base_urls((configured or "").strip()):
            if candidate and candidate not in collected:
                collected.append(candidate)
    return collected


def _solidset_login(
    client: httpx.Client,
    base_url: str,
    agent_resource_id: Optional[str] = None,
    agent_login_id: Optional[str] = None,
    host_header: Optional[str] = None,
    allow_sync_retry: bool = True,
) -> tuple[bool, str, str]:
    """Autentica la identidad global o la cuenta del recurso agente solicitado."""
    if agent_resource_id:
        print(
            f"🔐 Preparando LoginJson base={base_url} "
            f"agent_resource={agent_resource_id}",
            flush=True,
        )
        try:
            login = get_solidset_login_for_active_agent(
                agent_resource_id,
                preferred_login_id=agent_login_id,
            )
        except Exception as exc:
            print(f"❌ No se pudo leer SysLogin del agente: {exc}", flush=True)
            return False, "", ""
        # La consulta incluye el filtro SysResourceIA.active=true. Si no devuelve
        # fila, el recurso no puede actuar como agente ni utilizar otra identidad.
        if not login:
            print("❌ El agente no está activo o no tiene SysLogin asociado", flush=True)
            return False, "", ""
        username = str(login.get("Username") or "").strip()
        password = str(login.get("Password") or "")
        selected_login_id = str(login.get("IDLogin") or "").strip()
        if not username or not password:
            print("❌ SysLogin no contiene Username/Password utilizables", flush=True)
            return False, "", ""
        print(
            "🔐 SysLogin seleccionado "
            f"id_login={selected_login_id or '-'} username={username} "
            f"preferred={str(agent_login_id or '-').strip()}",
            flush=True,
        )

        login_payload = {
            "UserName": username,
            "Password": password,
            # SysLogin.Password ya contiene el HMAC generado por SolidSET.
            # LoginViewModel.PasswordEncrypted evita aplicar GenerateHMAC otra vez.
            "PasswordEncrypted": "true",
            # Coincide con LoginViewModel.TimezoneId. Resources es opcional y
            # no selecciona CurrentIDResource; se utiliza LastIDResource.
            "TimezoneId": settings.SOLIDSET_TIMEZONE_ID or "GMT Standard Time",
        }
        try:
            response = client.post(
                f"{base_url}/User/LoginJson",
                data=login_payload,
                headers=_solidset_action_headers(host_header),
            )
            print(
                f"🔐 LoginJson respondió HTTP {response.status_code} "
                f"base={base_url}",
                flush=True,
            )
            if response.status_code >= 400:
                print(
                    f"❌ LoginJson HTTP {response.status_code}: "
                    f"{response.text[:500]}",
                    flush=True,
                )
                return False, "", ""
            try:
                result = response.json()
            except Exception:
                result = None
            if isinstance(result, dict):
                success = result.get("Success", result.get("success"))
                error = result.get("Error", result.get("error"))
                if success is False or error:
                    print(
                        f"❌ LoginJson rechazado: {str(result)[:300]}",
                        flush=True,
                    )
                    if allow_sync_retry:
                        try:
                            from app.system.resource_ingest import ingest_solidset_logins

                            sync_result = ingest_solidset_logins()
                            print(
                                "🔄 SysLogin resincronizado después del rechazo "
                                f"updated={sync_result.get('updated', 0)} "
                                f"inserted={sync_result.get('inserted', 0)}",
                                flush=True,
                            )
                            return _solidset_login(
                                client,
                                base_url,
                                agent_resource_id=agent_resource_id,
                                agent_login_id=agent_login_id,
                                host_header=host_header,
                                allow_sync_retry=False,
                            )
                        except Exception as sync_exc:
                            print(
                                f"❌ No se pudo resincronizar SysLogin: {sync_exc}",
                                flush=True,
                            )
                    return False, "", ""
            return True, "/User/LoginJson", _solidset_extract_access_key(response)
        except Exception as exc:
            print(
                f"❌ Error conectando con LoginJson base={base_url}: {exc}",
                flush=True,
            )
            return False, "", ""

    if not settings.SOLIDSET_LOGIN_USERNAME and not settings.SOLIDSET_LOGIN_HASHPASS:
        return False, "", ""

    # API login (JSON)
    payload = {
        "username": settings.SOLIDSET_LOGIN_USERNAME or None,
        "pass": settings.SOLIDSET_LOGIN_PASSWORD or None,
        "hashPass": settings.SOLIDSET_LOGIN_HASHPASS or None,
        "accessKey": True,
        "generateAccessKey": False,
    }
    if settings.SOLIDSET_RESOURCE_ID:
        payload["resource"] = settings.SOLIDSET_RESOURCE_ID

    for endpoint in ["/api/User/LoginRaw", "/api/User/Login"]:
        try:
            resp = client.post(f"{base_url}{endpoint}", json=payload, headers=_solidset_action_headers())
            if resp.status_code < 400:
                access_key = _solidset_extract_access_key(resp)
                return True, endpoint, access_key
        except Exception:
            continue

    # Legacy login (form)
    legacy_payload = {
        "UserName": settings.SOLIDSET_LOGIN_USERNAME or "",
        "Password": settings.SOLIDSET_LOGIN_PASSWORD or "",
        "TimezoneID": settings.SOLIDSET_TIMEZONE_ID or "GMT Standard Time",
    }
    try:
        legacy = client.post(
            f"{base_url}/User/LoginJson",
            data=legacy_payload,
            headers=_solidset_action_headers(),
        )
        if legacy.status_code < 400:
            return True, "/User/LoginJson", ""
        return False, "", ""
    except Exception:
        return False, "", ""


def _solidset_authenticate_runtime(force: bool = False) -> tuple[bool, str]:
    current = _solidset_get_runtime_auth()
    if not force and current.get("base_url") and (
        current.get("access_key") or current.get("cookie_header")
    ):
        return True, str(current.get("base_url"))

    candidates = _solidset_get_all_base_candidates()
    if not candidates:
        return False, "no hay instancias activas registradas en SysSolidSETInstance"

    last_error = "sin detalle"
    for base in candidates:
        try:
            with httpx.Client(timeout=15.0, verify=settings.NOTIF_API_VERIFY_TLS, follow_redirects=True) as client:
                ok, endpoint, access_key = _solidset_login(client, base)
                if not ok:
                    last_error = f"No se pudo autenticar en {base}"
                    continue

                cookie_header = _solidset_client_cookie_header(client)
                workstation_cookie = (_solidset_action_headers().get("Cookie") or "").strip()
                merged_cookie = _solidset_merge_cookie_headers(workstation_cookie, cookie_header)
                _solidset_set_runtime_auth(
                    base_url=base,
                    cookie_header=merged_cookie,
                    access_key=access_key,
                    login_endpoint=endpoint,
                )
                return True, base
        except Exception as exc:
            last_error = str(exc)

    return False, last_error


def _solidset_authenticated_headers(base_headers: Optional[dict[str, str]] = None) -> dict[str, str]:
    headers = dict(base_headers or _solidset_action_headers())
    auth = _solidset_get_runtime_auth()
    runtime_cookie = (auth.get("cookie_header") or "").strip()
    base_cookie = (headers.get("Cookie") or "").strip()
    merged_cookie = _solidset_merge_cookie_headers(base_cookie, runtime_cookie)
    if merged_cookie:
        headers["Cookie"] = merged_cookie

    access_key = (auth.get("access_key") or "").strip()
    if access_key:
        headers["X-Access-Key"] = access_key
        headers["Authorization"] = f"Bearer {access_key}"
    return headers


def _solidset_request_authenticated(
    *,
    method: str,
    endpoint: str,
    params: Optional[dict[str, Any]] = None,
    json_payload: Optional[dict[str, Any]] = None,
    form_payload: Optional[dict[str, Any]] = None,
) -> tuple[Optional[httpx.Response], str, str]:
    ok, detail = _solidset_authenticate_runtime(force=False)
    if not ok:
        return None, "", f"No se pudo autenticar: {detail}"

    auth = _solidset_get_runtime_auth()
    base = (auth.get("base_url") or "").rstrip("/")
    if not base:
        return None, "", "No hay base URL autenticada"

    request_kwargs: dict[str, Any] = {
        "params": params or None,
        "headers": _solidset_authenticated_headers(),
    }
    if json_payload is not None:
        request_kwargs["json"] = json_payload
    if form_payload is not None:
        request_kwargs["data"] = form_payload

    target = endpoint if endpoint.startswith("/") else f"/{endpoint}"
    url = f"{base}{target}"

    try:
        with httpx.Client(timeout=20.0, verify=settings.NOTIF_API_VERIFY_TLS, follow_redirects=True) as client:
            response = client.request(method.upper(), url, **request_kwargs)
            if response.status_code in {401, 403}:
                # Reintento único con relogin forzado.
                ok_retry, detail_retry = _solidset_authenticate_runtime(force=True)
                if not ok_retry:
                    return None, base, f"autenticación expirada y no se pudo renovar: {detail_retry}"
                retry_kwargs = dict(request_kwargs)
                retry_kwargs["headers"] = _solidset_authenticated_headers()
                response = client.request(method.upper(), url, **retry_kwargs)
            return response, base, ""
    except Exception as exc:
        return None, base, str(exc)


def _solidset_request_as_agent(
    *,
    agent_resource_id: str,
    agent_login_id: Optional[str] = None,
    method: str,
    endpoint: str,
    form_payload: Optional[dict[str, Any]] = None,
    solidset_base_url: Optional[str] = None,
) -> tuple[Optional[httpx.Response], str, str]:
    """Autentica y ejecuta una petición usando la cuenta del recurso agente."""
    target = endpoint if endpoint.startswith("/") else f"/{endpoint}"
    last_error = "el recurso no está activo, no tiene SysLogin o SolidSET rechazó LoginJson"
    bases = _solidset_candidate_base_urls(str(solidset_base_url or "").strip())
    if not bases:
        return None, "", "la petición no contiene una instancia SolidSET registrada"
    for base in bases:
        configured_url = urlparse(str(solidset_base_url or "").strip())
        effective_url = urlparse(base)
        logical_host = ""
        if (
            (configured_url.hostname or "").lower() in {"localhost", "127.0.0.1"}
            and (effective_url.hostname or "").lower() == "host.docker.internal"
        ):
            # IIS puede escuchar físicamente en el host Docker pero aceptar
            # únicamente el binding lógico localhost:puerto.
            logical_host = configured_url.netloc
        print(
            f"🌐 Intentando envío autenticado del agente base={base} "
            f"host_header={logical_host or '-'}",
            flush=True,
        )
        try:
            with httpx.Client(
                timeout=20.0,
                verify=settings.NOTIF_API_VERIFY_TLS,
                follow_redirects=True,
            ) as client:
                authenticated, _, access_key = _solidset_login(
                    client,
                    base,
                    agent_resource_id=agent_resource_id,
                    agent_login_id=agent_login_id,
                    host_header=logical_host or None,
                )
                if not authenticated:
                    print(f"⚠️ Login del agente falló en {base}", flush=True)
                    continue

                cookie_header = _solidset_merge_cookie_headers(
                    _solidset_action_headers(logical_host).get("Cookie", ""),
                    _solidset_client_cookie_header(client),
                )
                headers = _solidset_action_headers(logical_host)
                if cookie_header:
                    headers["Cookie"] = cookie_header
                if access_key:
                    headers["X-Access-Key"] = access_key
                    headers["Authorization"] = f"Bearer {access_key}"
                response = client.request(
                    method.upper(),
                    f"{base}{target}",
                    headers=headers,
                    data=form_payload,
                )
                print(
                    f"📨 Chat/SendMessageForm respondió HTTP "
                    f"{response.status_code} base={base}",
                    flush=True,
                )
                return response, base, ""
        except Exception as exc:
            last_error = str(exc)
            print(f"❌ Error enviando mediante {base}: {exc}", flush=True)
    return None, "", last_error


def _normalize_document_lines(content: str) -> list[str]:
    text = (content or "").replace("\r\n", "\n").replace("\r", "\n")
    lines = [line.rstrip() for line in text.split("\n")]
    # Compactar bloques vacíos consecutivos para evitar saltos excesivos.
    normalized: list[str] = []
    previous_blank = False
    for line in lines:
        is_blank = not line.strip()
        if is_blank and previous_blank:
            continue
        normalized.append(line)
        previous_blank = is_blank
    return normalized


def _infer_document_kind(title: str, content: str) -> str:
    probe = f"{title} {content}".lower()
    if any(word in probe for word in ["acta", "meeting", "minuta"]):
        return "Acta"
    if any(word in probe for word in ["resumen", "summary"]):
        return "Resumen"
    if any(word in probe for word in ["canal", "incidencia", "diagnost", "informe", "report"]):
        return "Informe"
    return "Documento"


def _extract_cte_names(query: str) -> set[str]:
    cte_names = set()
    for match in re.finditer(r"(?:WITH|,)\s*([A-Za-z_][A-Za-z0-9_]*)\s+AS\s*\(", query, flags=re.IGNORECASE):
        cte_names.add(match.group(1).lower())
    return cte_names


def _extract_table_references(query: str) -> set[str]:
    refs = set()
    cte_names = _extract_cte_names(query)
    for match in re.finditer(r"\b(?:FROM|JOIN)\s+([A-Za-z0-9_\[\]\.]+)", query, flags=re.IGNORECASE):
        token = match.group(1).strip()
        if token.startswith("("):
            continue
        clean = token.replace("[", "").replace("]", "")
        parts = [p for p in clean.split(".") if p]
        if not parts:
            continue
        if len(parts) == 1:
            schema = "dbo"
            table = parts[0]
        else:
            schema = parts[-2]
            table = parts[-1]
        if table.lower() in cte_names:
            continue
        refs.add(f"{schema.lower()}.{table.lower()}")
    return refs


def _load_real_tables(cursor) -> set[str]:
    cursor.execute(
        """
        SELECT TABLE_SCHEMA, TABLE_NAME
        FROM INFORMATION_SCHEMA.TABLES
        WHERE TABLE_TYPE = 'BASE TABLE'
        """
    )
    rows = cursor.fetchall() or []
    return {f"{str(r['TABLE_SCHEMA']).lower()}.{str(r['TABLE_NAME']).lower()}" for r in rows}


# ---------------------------------------------------------------------------
# 1. TOOL: SQL Server Query (MEJORADA con validación de filas)
# ---------------------------------------------------------------------------
@tool
def _legacy_google_web_search(query: str) -> str:
    """
    REALIZA UNA BÚSQUEDA EN INTERNET USANDO GOOGLE.

    CUÁNDO USARLA:
    - Cuando el usuario pregunte por información que no esté en la base de datos o en la documentación técnica.
    - Cuando el usuario pregunte por noticias, eventos actuales o temas generales de conocimiento.
    - Cuando una consulta sobre un error, componente o término técnico no devuelva resultados en los sistemas internos.
    
    EJEMPLOS:
    - "¿Cuál es el precio actual del acero 1045?"
    - "Busca información sobre el error 'F3001' en un controlador Fanuc."
    - "Noticias sobre la industria metalmecánica."
    """
    # Esta es una implementación de marcador de posición.
    # El usuario debe reemplazar esto con una llamada a una API de búsqueda web real.
    try:
        # Aquí iría la lógica de búsqueda real.
        # Por ejemplo, usando una biblioteca como 'googlesearch-python':
        # from googlesearch import search
        # search_results = []
        # for result_url in search(query, num=3, stop=3, pause=1):
        #     search_results.append(result_url)
        # if not search_results:
        #     return f"No se encontraron resultados para '{query}'."
        # return f"Resultados de búsqueda para '{query}':\n" + "\n".join(search_results)
        
        print(f"--- SIMULANDO BÚSQUEDA WEB PARA: '{query}' ---")
        return f"Resultados de búsqueda para '{query}':\n- https://es.wikipedia.org/wiki/Acero\n- https://www.ejemplo-proveedor.com/acero-1045\n- https://www.machining-forums.com/error-f3001-fanuc"

    except Exception as e:
        return f"Error durante la búsqueda web: {e}"


def _store_web_search_knowledge(query: str, results: list[dict[str, str]]) -> bool:
    """Index web results with deterministic IDs and full provenance."""
    client = QdrantClient(url=settings.VECTOR_DB_URL)
    embeddings = OllamaEmbeddings(base_url=settings.OLLAMA_BASE_URL, model=settings.EMBEDDING_MODEL_NAME)
    collections = [c.name for c in client.get_collections().collections]
    probe_vector = None
    if settings.VECTOR_COLLECTION_NAME not in collections:
        probe_vector = embeddings.embed_query(query)
        client.create_collection(
            collection_name=settings.VECTOR_COLLECTION_NAME,
            vectors_config=VectorParams(size=len(probe_vector), distance=Distance.COSINE),
        )

    learned_at = datetime.now().astimezone().isoformat()
    points = []
    for index, result in enumerate(results):
        content = (
            f"Consulta web: {query}\nTítulo: {result['title']}\n"
            f"Resumen: {result['snippet']}\nFuente: {result['url']}"
        )
        vector = probe_vector if index == 0 and probe_vector is not None else embeddings.embed_query(content)
        digest = hashlib.md5(f"web:{result['url']}:{content}".encode("utf-8")).hexdigest()
        points.append(PointStruct(
            id=str(uuid.UUID(digest)),
            vector=vector,
            payload={
                "page_content": content,
                "category": "web_research",
                "source": "web_search",
                "source_url": result["url"],
                "source_title": result["title"],
                "search_query": query,
                "external_unverified": True,
                "learned_at": learned_at,
            },
        ))
    client.upsert(collection_name=settings.VECTOR_COLLECTION_NAME, points=points)
    return True


def _schedule_web_search_learning(query: str, results: list[dict[str, str]]) -> None:
    """Difiere la indexación para que embeddings y chat no compitan en Ollama."""
    def _learn() -> None:
        try:
            _store_web_search_knowledge(query, results)
            print(f"🧠 Resultados web indexados en background; query={query[:80]!r}")
        except Exception as exc:
            print(f"⚠️ No se pudieron indexar resultados web en background: {exc}")

    timer = threading.Timer(120.0, _learn)
    timer.daemon = True
    timer.start()


@tool
def google_web_search(query: str) -> str:
    """Search the public web and persist results with source provenance for later retrieval."""
    started_at = perf_counter()
    try:
        clean_query = " ".join((query or "").split())
        if not clean_query:
            return "Error: la consulta de búsqueda no puede estar vacía."
        if not settings.WEB_SEARCH_ENABLED:
            return "La búsqueda web está desactivada por configuración."

        from ddgs import DDGS

        with DDGS(timeout=settings.WEB_SEARCH_TIMEOUT_SECONDS) as client:
            raw_results = list(client.text(
                clean_query,
                region=settings.WEB_SEARCH_REGION,
                safesearch=settings.WEB_SEARCH_SAFESEARCH,
                max_results=settings.WEB_SEARCH_MAX_RESULTS,
            ))

        results = []
        seen_urls = set()
        for item in raw_results:
            url = str(item.get("href") or item.get("url") or "").strip()
            if not url or url in seen_urls:
                continue
            seen_urls.add(url)
            results.append({
                "title": str(item.get("title") or "Sin título").strip()[:300],
                "snippet": str(item.get("body") or item.get("snippet") or "").strip()[:1200],
                "url": url[:2000],
            })
        if not results:
            return f"No se encontraron resultados web para '{clean_query}'."

        learned = False
        learning_scheduled = False
        if settings.WEB_SEARCH_AUTO_LEARN:
            # Indexar puede requerir varios embeddings de Ollama. Se desacopla de la
            # respuesta para no añadir minutos de espera al usuario.
            _schedule_web_search_learning(clean_query, results)
            learning_scheduled = True

        payload = {
            "query": clean_query,
            "source_type": "web_search",
            "external_unverified": True,
            "learned": learned,
            "learning_scheduled": learning_scheduled,
            "results": results,
        }
        elapsed = perf_counter() - started_at
        print(
            f"🌐 Búsqueda web completada en {elapsed:.2f}s; "
            f"resultados={len(results)} aprendizaje_background={learning_scheduled}"
        )
        return json.dumps(payload, ensure_ascii=False)
    except Exception as exc:
        return f"Error durante la búsqueda web: {str(exc)[:300]}"


@tool
def query_sql_server(query: str) -> str:
    """
    EJECUTA CONSULTAS SELECT EN SQL SERVER.
    
    CUÁNDO USARLA:
    - Cuando el usuario pida datos de clientes (dbo.Account)
    - Cuando el usuario pida historial de actividades (dbo.Activity)
    - Cuando el usuario pida información de máquinas/activos (dbo.Asset)
    - Cuando el usuario pregunte por saldos, deudas o inventarios
    
    CUÁNDO NO USARLA:
    - NO la uses para explorar la estructura de tablas (usa get_db_schema)
    - NO la uses para saludos o conversación general
    
    ADVERTENCIA DE SEGURIDAD:
    - Siempre usa filtros WHERE para evitar consultas masivas
    - Deriva el filtro desde la petición del usuario cuando mencione un nombre, alias, estado o patrón.
    - No preguntes qué tabla usar si el esquema conocido ya identifica la tabla y columnas.
    - Los agregados COUNT con WHERE son consultas acotadas y no requieren confirmación.
    """
    server = settings.SQL_SERVER_HOST
    user = settings.SQL_SERVER_USER
    password = settings.SQL_SERVER_PASSWORD
    database = settings.SQL_SERVER_DB

    clean_query = query.strip()
    if not clean_query.upper().startswith("SELECT") and not clean_query.upper().startswith("WITH"):
        return "Error de seguridad: Solo se permiten consultas de lectura (SELECT / WITH)."

    forbidden_keywords = ["DELETE", "INSERT", "UPDATE", "DROP", "ALTER", "TRUNCATE", "EXEC", "EXECUTE"]
    if any(kw in clean_query.upper() for kw in forbidden_keywords):
        return "Error de seguridad: La consulta contiene comandos no permitidos."

    try:
        conn = pymssql.connect(
            **settings.sql_server_connection_options(),
            user=user,
            password=password,
            database=database,
            timeout=5,
        )
        cursor = conn.cursor(as_dict=True)

        referenced_tables = _extract_table_references(clean_query)
        real_tables = _load_real_tables(cursor)
        unknown_tables = sorted(t for t in referenced_tables if t not in real_tables)
        if unknown_tables:
            conn.close()
            unknown = ", ".join(unknown_tables)
            sample = ", ".join(sorted(list(real_tables))[:12])
            return (
                "Error de esquema: la consulta usa tablas que no existen en esta base de datos: "
                f"{unknown}. "
                "Consulta primero el esquema real con get_db_schema antes de reintentar. "
                f"Ejemplos de tablas reales: {sample}."
            )

        cursor.execute(clean_query)
        rows = cursor.fetchall()
        conn.close()

        if not rows:
            return "La consulta se ejecutó correctamente pero no devolvió resultados."

        # 🚨 NUEVA VALIDACIÓN: Si son más de 50 filas, advertir
        if len(rows) > 50:
            return f"⚠️ ADVERTENCIA: La consulta devolvió {len(rows)} filas. Mostrando solo las primeras 15.\n\n{json.dumps(rows[:15])}"

        return json.dumps(rows[:15])

    except pymssql.Error as db_err:
        return f"Error SQL Server: {str(db_err)}. Ajusta los campos/tablas y vuelve a intentar."
    except Exception as e:
        return f"Error al conectar o consultar SQL Server: {str(e)}"


@tool
def get_db_schema(table_name: Optional[str] = None) -> str:
    """
    EXPLORA LA ESTRUCTURA DE LA BASE DE DATOS.
    
    CUÁNDO USARLA (SOLO PARA EXPLORACIÓN):
    - El usuario pregunta "¿Qué tablas hay en la base de datos?"
    - El usuario pregunta "¿Qué columnas tiene la tabla X?"
    - El usuario dice "No sé qué tabla contiene los datos de clientes"
    
    CUÁNDO NO USARLA:
    - NO la uses para obtener datos reales de negocio (usa query_sql_server)
    - NO la uses si el usuario ya sabe qué tabla consultar
    """
    server = settings.SQL_SERVER_HOST
    user = settings.SQL_SERVER_USER
    password = settings.SQL_SERVER_PASSWORD
    database = settings.SQL_SERVER_DB

    try:
        conn = pymssql.connect(
            **settings.sql_server_connection_options(),
            user=user,
            password=password,
            database=database,
            timeout=5,
        )
        cursor = conn.cursor(as_dict=True)

        if table_name:
            query = """
                SELECT COLUMN_NAME, DATA_TYPE, IS_NULLABLE
                FROM INFORMATION_SCHEMA.COLUMNS
                WHERE TABLE_NAME = %s
            """
            cursor.execute(query, (table_name,))
            rows = cursor.fetchall()
            conn.close()
            if not rows:
                return f"No se encontró la tabla '{table_name}'."
            # Formatear bonito
            resultado = f"📋 ESTRUCTURA DE LA TABLA '{table_name}':\n\n"
            for row in rows:
                resultado += f"  • {row['COLUMN_NAME']} ({row['DATA_TYPE']}) - {'Puede ser NULL' if row['IS_NULLABLE'] == 'YES' else 'NO NULL'}\n"
            return resultado
        else:
            query = """
                SELECT TABLE_NAME 
                FROM INFORMATION_SCHEMA.TABLES 
                WHERE TABLE_TYPE = 'BASE TABLE'
                ORDER BY TABLE_NAME
            """
            cursor.execute(query)
            rows = cursor.fetchall()
            conn.close()
            tables = [r['TABLE_NAME'] for r in rows]
            return f"📊 TABLAS DISPONIBLES ({len(tables)} en total):\n\n" + "\n".join([f"  • {t}" for t in tables[:30]])

    except Exception as e:
        return f"Error al consultar el esquema: {str(e)}"


# ---------------------------------------------------------------------------
# 🚨 NUEVA TOOL: Confirmación de consultas pesadas (Human-in-the-loop)
# ---------------------------------------------------------------------------
@tool
def confirm_large_operation(operation_type: str, description: str, estimated_impact: str) -> str:
    """
    SOLICITA CONFIRMACIÓN DEL USUARIO ANTES DE OPERACIONES CRÍTICAS.
    
    CUÁNDO USARLA (SIEMPRE antes de):
    - Ejecutar un SELECT sin filtros WHERE (puede devolver miles de filas)
    - Enviar correos o mensajes automáticos
    - Realizar cambios en configuración de máquinas
    - Cualquier acción que el usuario NO haya pedido explícitamente
    
    CÓMO USARLA:
    1. Describe la operación que quieres hacer
    2. Espera la respuesta del usuario (Sí/No)
    3. Si dice "Sí", ejecuta la acción; si dice "No", cancela
    """
    return f"⚠️ SOLICITUD DE CONFIRMACIÓN:\n\nOperación: {operation_type}\nDescripción: {description}\nImpacto estimado: {estimated_impact}\n\n❓ ¿Confirmas que quieres realizar esta operación? (Responde 'Sí' o 'No')"


# ---------------------------------------------------------------------------
# 2. TOOL: Consumir API Externa (MEJORADA)
# ---------------------------------------------------------------------------
@tool
def fetch_external_api(endpoint_url: str, method: str = "GET", payload: Optional[Union[dict, str]] = None,) -> str:
    """
    CONSULTA APIs EXTERNAS (HTTP/GraphQL).
    
    CUÁNDO USARLA:
    - El usuario comparte una URL y dice "consulta esto"
    - El usuario pide datos de un sistema externo (clima, cotizaciones, etc.)
    
    PRECAUCIÓN:
    - NUNCA la uses para APIs internas de la planta (usa query_sql_server)
    - Siempre valida que la URL sea segura
    """
    try:
        parsed_payload = None
        if payload:
            if isinstance(payload, str):
                try:
                    parsed_payload = json.loads(payload)
                except json.JSONDecodeError:
                    clean_str = payload.replace("'", '"')
                    parsed_payload = json.loads(clean_str)
            else:
                parsed_payload = payload

        with httpx.Client(timeout=15.0) as client:
            if method.upper() == "POST":
                response = client.post(endpoint_url, json=parsed_payload or {})
            else:
                response = client.get(endpoint_url)
                
            response.raise_for_status()
            return str(response.json())
            
    except httpx.HTTPStatusError as exc:
        return f"Error HTTP {exc.response.status_code}: {exc.response.text}"
    except Exception as e:
        return f"Error de conexión con la API: {str(e)}"


@tool
def solidset_send_chat_message(
    canal_id: Optional[str],
    mensaje: str,
    importance: Union[int, str] = 1,
    kind: int = 7,
    visibility_level: Union[int, str] = 1,
    confirm: bool = False,
    recurso_id: Optional[str] = None,
    recurso_login_id: Optional[str] = None,
    meeting_id: Optional[str] = None,
    meeting_code: Optional[str] = None,
    meeting_mirror_general: bool = False,
    generated_by_ia: bool = False,
    agent_resource_id: Optional[str] = None,
    agent_identity_id: Optional[str] = None,
    agent_chat_resource_name: Optional[str] = None,
    agent_chat_login_id: Optional[str] = None,
    human_chat_resource_name: Optional[str] = None,
    solidset_base_url: Optional[str] = None,
) -> str:
    """
    ENVÍA UN MENSAJE REAL AL CHAT/CANAL DE SOLIDSET COMO USUARIO AUTENTICADO.

    Reglas:
    - Solo se ejecuta cuando SOLIDSET_USER_ACTIONS_ENABLED=true.
    - Requiere confirm=true para evitar envíos accidentales.
    - Usa credenciales SOLIDSET_LOGIN_* del .env.
    """
    print(
        "📨 Entrando en solidset_send_chat_message "
        f"base={solidset_base_url or '-'} channel={canal_id or '-'} "
        f"resource={recurso_id or '-'} meeting={meeting_id or '-'}",
        flush=True,
    )
    if not settings.SOLIDSET_USER_ACTIONS_ENABLED:
        return (
            "Acción bloqueada: habilita SOLIDSET_USER_ACTIONS_ENABLED=true en .env "
            "para permitir operaciones de escritura en SOLIDSET."
        )

    if not confirm:
        return (
            "⚠️ Confirmación requerida. Repite la acción con confirm=true "
            "si deseas enviar el mensaje al canal real de SOLIDSET."
        )

    channel = (canal_id or "").strip()
    resource = (recurso_id or "").strip()
    resource_login = (recurso_login_id or "").strip()
    text = (mensaje or "").strip()
    if not channel and not resource:
        return "Error: canal_id o recurso_id es obligatorio."
    if not text:
        return "Error: mensaje está vacío."

    visibility_names = {
        "public": 0,
        "normal": 1,
        "confidential": 2,
        "private": 3,
    }
    if isinstance(visibility_level, str):
        visibility_normalized = visibility_names.get(visibility_level.strip().lower())
        if visibility_normalized is None and visibility_level.strip().isdigit():
            visibility_normalized = int(visibility_level.strip())
    else:
        visibility_normalized = int(visibility_level)
    if visibility_normalized not in {0, 1, 2, 3}:
        visibility_normalized = 1

    importance_names = {
        "low": 0,
        "normal": 1,
        "high": 2,
        "urgent": 3,
    }
    if isinstance(importance, str):
        importance_normalized = importance_names.get(importance.strip().lower())
        if importance_normalized is None and importance.strip().isdigit():
            importance_normalized = int(importance.strip())
    else:
        importance_normalized = int(importance)
    if importance_normalized not in {0, 1, 2, 3}:
        importance_normalized = 1

    form_payload = {
        "Importance": importance_normalized,
        "Kind": int(kind),
        "VisibilityLevel": visibility_normalized,
        "RawMessage": text,
    }
    if channel:
        form_payload["Destiny.WorkRoom"] = channel
    if channel and resource:
        # Inversión explícita del Chat.destiny original:
        # solicitud humano(type=1) -> IA(type=2), respuesta IA -> humano.
        # SolidSET obtiene el origen IA desde la sesión y las marcas Info; aquí
        # el humano queda como destinatario dirigido Kind/Type=2.
        if resource_login:
            form_payload["Destiny.Dests[0].Login"] = resource_login
        form_payload["Destiny.Dests[0].Resource"] = resource
        form_payload["Destiny.Dests[0].Kind"] = 2
        # Compatibilidad con la nueva variante del contrato de SolidSET, que
        # identifica explícitamente las intervenciones de IA mediante type=2.
        form_payload["Destiny.Dests[0].Type"] = 2
    elif resource:
        form_payload["Destiny.Resource"] = resource
        if resource_login:
            form_payload["Destiny.Login"] = resource_login
    requested_meeting = str(meeting_id or "").strip()
    meeting_label = str(meeting_code or "").strip()
    meeting = _resolve_solidset_meeting_id(
        requested_meeting,
        channel,
        meeting_label,
    ) if requested_meeting else None
    if requested_meeting and not meeting:
        print(
            "⚠️ Meeting descartado para evitar FK inválida: "
            f"meeting={requested_meeting} code={meeting_label or '-'} "
            f"channel={channel or '-'}"
        )
    if meeting:
        # SolidSET vincula el chat al meeting a través de ExtraData. El workroom
        # permanece únicamente como ruta técnica; no se marca como espejo general.
        meeting_extra = {"meeting_id": meeting}
        if meeting_label:
            meeting_extra["meeting_code"] = meeting_label
        form_payload["ExtraData"] = json.dumps(meeting_extra, separators=(",", ":"))
        form_payload["Info[meeting_id]"] = meeting
        if meeting_label:
            form_payload["Info[meeting_code]"] = meeting_label
    elif meeting_mirror_general:
        # Compatibilidad con el modo espejo antiguo cuando no existe meeting_id.
        form_payload["Info[meeting_mirror_general]"] = "1"
    if generated_by_ia:
        form_payload["Info[generated_by_ia]"] = "1"
        sender_agent_resource_id = str(
            agent_identity_id or agent_resource_id or ""
        ).strip()
        selected_agent_resource_id = str(agent_resource_id or "").strip()
        if sender_agent_resource_id:
            form_payload["Info[agent_resource_id]"] = sender_agent_resource_id
        if agent_identity_id:
            visual_agent_id = str(agent_identity_id).strip()
            # Identidad lógica de la autorrespuesta. SolidSET autentica y persiste
            # el emisor real desde la sesión; el cliente debe usar esta marca para
            # no clasificar el mensaje del agente como FromSelf.
            form_payload["Info[id_agent_ia]"] = visual_agent_id
            form_payload["Info[agent_id]"] = visual_agent_id
            form_payload["IDAgentIA"] = visual_agent_id
        if sender_agent_resource_id and selected_agent_resource_id and channel and resource:
            # La UI de SolidSET pinta From/To desde Chat, no solo desde Destiny.
            # Para una autorrespuesta se invierten explícitamente los roles:
            # agente lógico type=1 (origen) -> humano type=2 (destino).
            # Contrato requerido por el receptor SolidSET: estos dos campos se
            # envían invertidos respecto al FrameworkMessage entrante.
            sender_login_id = str(agent_chat_login_id or resource_login or "").strip()
            if sender_login_id:
                form_payload["Chat.IDSenderResource"] = sender_login_id
            form_payload["Chat.IDSender"] = sender_agent_resource_id
            form_payload["Chat.IDWorkRoom"] = channel
            if meeting:
                form_payload["Chat.IDMeeting"] = meeting
            form_payload["Chat.RawMessage"] = text
            form_payload["Chat.Kind"] = 60 if meeting else 0
            if agent_chat_login_id:
                form_payload["Chat.Destiny[0].IDLogin"] = str(agent_chat_login_id).strip()
            # El From siempre utiliza el IDAgentResource verificado contra
            # dbo.SysResource2Agent. El recurso humano queda exclusivamente
            # como destinatario Type=2 en Destiny[1].
            form_payload["Chat.Destiny[0].IDResource"] = sender_agent_resource_id
            form_payload["Chat.Destiny[0].ResourceName"] = (
                str(agent_chat_resource_name or "").strip() or "Agente IA"
            )
            form_payload["Chat.Destiny[0].TalkWithAgent"] = "true"
            form_payload["Chat.Destiny[0].Type"] = 1
            form_payload["Chat.Destiny[0].IDChannel"] = channel
            form_payload["Chat.Destiny[0].Sequence"] = 0
            if resource_login:
                form_payload["Chat.Destiny[1].IDLogin"] = resource_login
            form_payload["Chat.Destiny[1].IDResource"] = resource
            if human_chat_resource_name:
                form_payload["Chat.Destiny[1].ResourceName"] = str(
                    human_chat_resource_name
                ).strip()
            form_payload["Chat.Destiny[1].Type"] = 2
            form_payload["Chat.Destiny[1].IDChannel"] = channel
            form_payload["Chat.Destiny[1].Sequence"] = 1
    print(form_payload, flush=True)
    request_sender = _solidset_request_as_agent if agent_resource_id else _solidset_request_authenticated    
    request_args: dict[str, Any] = {
        "method": "POST",
        "endpoint": "/Chat/SendMessageForm",
        # SendMessageForm espera datos de formulario. Usar ``params`` colocaba
        # RawMessage en la URL y hacía que IIS devolviera 404.15 para respuestas
        # extensas (por ejemplo, análisis construidos desde cientos de mensajes).
        "form_payload": form_payload,
    }
    if agent_resource_id:
        request_args["agent_resource_id"] = str(agent_resource_id).strip()
        request_args["agent_login_id"] = str(agent_chat_login_id or "").strip() or None
        request_args["solidset_base_url"] = str(solidset_base_url or "").strip()
    response, base, error = request_sender(**request_args)
    if response is None:
        return f"Error enviando mensaje a SOLIDSET: {error or 'sin detalle'}"
    if response.status_code >= 400:
        return f"Error enviando mensaje a SOLIDSET: HTTP {response.status_code} -> {response.text[:220]}"
    try:
        result_payload = response.json()
    except Exception:
        result_payload = None
    if isinstance(result_payload, dict):
        success = result_payload.get("Success", result_payload.get("success"))
        result_code = result_payload.get("Result", result_payload.get("result"))
        api_error = result_payload.get("Error", result_payload.get("error"))
        if success is False or (isinstance(result_code, int) and result_code != 0) or api_error:
            return (
                "Error enviando mensaje a SOLIDSET: rechazo funcional "
                f"HTTP {response.status_code} -> {str(result_payload)[:220]}"
            )
    return (
        f"✅ Mensaje enviado a "
        f"{'recurso ' + resource + ' dentro del canal ' + channel if channel and resource else ('canal ' + channel if channel else 'chat del recurso ' + resource)}. "
        f"Endpoint: {base}/Chat/SendMessageForm"
    )


@tool
def solidset_update_reaction(
    id_chat: int,
    reaction: str,
    id_user: Optional[str] = None,
    confirm: bool = False,
) -> str:
    """
    REGISTRA UNA REACCIÓN EN UN MENSAJE/DUDA DEL CHAT DE SOLIDSET.

    Reglas:
    - Solo se ejecuta cuando SOLIDSET_USER_ACTIONS_ENABLED=true.
    - Requiere confirm=true para evitar cambios accidentales.
    - Usa credenciales SOLIDSET_LOGIN_* del .env.
    """
    if not settings.SOLIDSET_USER_ACTIONS_ENABLED:
        return (
            "Acción bloqueada: habilita SOLIDSET_USER_ACTIONS_ENABLED=true en .env "
            "para permitir operaciones de escritura en SOLIDSET."
        )

    if not confirm:
        return (
            "⚠️ Confirmación requerida. Repite la acción con confirm=true "
            "si deseas registrar la reacción en SOLIDSET."
        )

    if not isinstance(id_chat, int) or id_chat <= 0:
        return "Error: id_chat debe ser un entero positivo."

    reaction_value = (reaction or "").strip()
    if not reaction_value:
        return "Error: reaction está vacía."

    payload = {
        "IDChat": id_chat,
        "Reaction": reaction_value,
    }
    if id_user:
        payload["IDUser"] = id_user.strip()
    response, base, error = _solidset_request_authenticated(
        method="POST",
        endpoint="/chat/update-reaction",
        json_payload=payload,
    )
    if response is None:
        return f"Error registrando reacción en SOLIDSET: {error or 'sin detalle'}"
    if response.status_code >= 400:
        return f"Error registrando reacción en SOLIDSET: HTTP {response.status_code} -> {response.text[:220]}"
    return (
        f"✅ Reacción '{reaction_value}' registrada en chat {id_chat}. "
        f"Endpoint: {base}/chat/update-reaction"
    )


@tool
def solidset_authenticate(force_login: bool = False) -> str:
    """
    AUTENTICA EL AGENTE CONTRA SOLIDSET Y GUARDA LA SESIÓN (cookies/access key).

    Reglas:
    - Usa credenciales SOLIDSET_LOGIN_* del .env.
    - Si ya existe sesión, la reutiliza salvo force_login=true.
    """
    ok, detail = _solidset_authenticate_runtime(force=bool(force_login))
    if not ok:
        return f"Error de autenticación SOLIDSET: {detail}"

    auth = _solidset_get_runtime_auth()
    has_cookie = bool((auth.get("cookie_header") or "").strip())
    has_key = bool((auth.get("access_key") or "").strip())
    return (
        "✅ Autenticación SOLIDSET activa. "
        f"Base: {auth.get('base_url')}. "
        f"Login: {auth.get('login_endpoint') or 'desconocido'}. "
        f"Cookie: {'sí' if has_cookie else 'no'}. "
        f"AccessKey: {'sí' if has_key else 'no'}."
    )


@tool
def solidset_logout(confirm: bool = False) -> str:
    """
    CIERRA LA SESIÓN ACTUAL DE SOLIDSET EN EL AGENTE Y limpia credenciales en memoria.
    """
    if not confirm:
        return "⚠️ Confirmación requerida. Repite la acción con confirm=true para cerrar sesión en SOLIDSET."

    auth = _solidset_get_runtime_auth()
    base = (auth.get("base_url") or "").rstrip("/")
    if not base:
        _solidset_clear_runtime_auth()
        return "No había una sesión activa en memoria."

    try:
        with httpx.Client(timeout=15.0, verify=settings.NOTIF_API_VERIFY_TLS, follow_redirects=True) as client:
            response = client.post(
                f"{base}/User/LogOffJson",
                headers=_solidset_authenticated_headers(),
            )
            _solidset_clear_runtime_auth()
            if response.status_code < 400:
                return f"✅ Sesión cerrada correctamente en {base}/User/LogOffJson"
            return (
                "Sesión local limpiada, pero el backend devolvió "
                f"HTTP {response.status_code} al cerrar sesión."
            )
    except Exception as exc:
        _solidset_clear_runtime_auth()
        return f"Sesión local limpiada, pero falló el logoff remoto: {exc}"


@tool
def solidset_request(
    endpoint: str,
    method: str = "GET",
    query_json: Optional[str] = None,
    body_json: Optional[str] = None,
    form_json: Optional[str] = None,
    confirm: bool = False,
) -> str:
    """
    EJECUTA UNA LLAMADA AUTENTICADA A CUALQUIER ENDPOINT DE SOLIDSET.

    Reglas:
    - Autentica primero (login) y reutiliza sesión.
    - Para métodos con escritura (POST/PUT/PATCH/DELETE), exige confirm=true.
    """
    target = (endpoint or "").strip()
    if not target:
        return "Error: endpoint es obligatorio."
    if not target.startswith("/"):
        target = f"/{target}"

    verb = (method or "GET").strip().upper()
    allowed = {"GET", "POST", "PUT", "PATCH", "DELETE"}
    if verb not in allowed:
        return f"Error: método '{verb}' no soportado. Usa uno de: {', '.join(sorted(allowed))}."

    if verb in {"POST", "PUT", "PATCH", "DELETE"} and not confirm:
        return (
            "⚠️ Confirmación requerida. Repite la acción con confirm=true "
            "para ejecutar operaciones con escritura en SOLIDSET."
        )

    def _parse_optional_json(raw: Optional[str], field_name: str) -> Optional[dict[str, Any]]:
        if raw is None or not str(raw).strip():
            return None
        try:
            parsed = json.loads(raw)
        except Exception as exc:
            raise ValueError(f"{field_name} no es JSON válido: {exc}") from exc
        if not isinstance(parsed, dict):
            raise ValueError(f"{field_name} debe ser un objeto JSON (clave/valor).")
        return parsed

    try:
        query_data = _parse_optional_json(query_json, "query_json")
        body_data = _parse_optional_json(body_json, "body_json")
        form_data = _parse_optional_json(form_json, "form_json")
    except ValueError as exc:
        return f"Error: {exc}"

    if body_data is not None and form_data is not None:
        return "Error: usa body_json o form_json, pero no ambos a la vez."

    response, base, error = _solidset_request_authenticated(
        method=verb,
        endpoint=target,
        params=query_data,
        json_payload=body_data,
        form_payload=form_data,
    )
    if response is None:
        return f"Error llamando endpoint SOLIDSET: {error or 'sin detalle'}"

    content_type = (response.headers.get("content-type") or "").lower()
    payload_preview: str
    if "json" in content_type:
        try:
            payload_preview = json.dumps(response.json(), ensure_ascii=False)[:3500]
        except Exception:
            payload_preview = response.text[:3500]
    else:
        payload_preview = response.text[:3500]

    return (
        f"status={response.status_code}; method={verb}; url={base}{target}; "
        f"body={payload_preview}"
    )


def _solidset_preview_response(response: httpx.Response, max_chars: int = 3500) -> str:
    content_type = (response.headers.get("content-type") or "").lower()
    if "json" in content_type:
        try:
            return json.dumps(response.json(), ensure_ascii=False)[:max_chars]
        except Exception:
            return response.text[:max_chars]
    return response.text[:max_chars]


def _solidset_add_indexed_params(params: dict[str, Any], key: str, values: list[str]) -> dict[str, Any]:
    for idx, value in enumerate(values):
        params[f"{key}[{idx}]"] = value
    return params


@tool
def solidset_chat_get_targets(
    mode: int = 1,
    include_user_read_pointers: bool = False,
    include_tabs: bool = False,
) -> str:
    """
    LEE LOS DESTINOS/CHATS DISPONIBLES DEL USUARIO AUTENTICADO EN SOLIDSET.
    Endpoint: GET /Chat/GetAllChatTargets
    """
    params = {
        "mode": int(mode),
        "includeUserReadPointers": str(bool(include_user_read_pointers)).lower(),
        "includeTabs": str(bool(include_tabs)).lower(),
    }
    response, base, error = _solidset_request_authenticated(
        method="GET",
        endpoint="/Chat/GetAllChatTargets",
        params=params,
    )
    if response is None:
        return f"Error consultando destinos de chat: {error or 'sin detalle'}"
    return (
        f"status={response.status_code}; method=GET; "
        f"url={base}/Chat/GetAllChatTargets; body={_solidset_preview_response(response)}"
    )


@tool
def solidset_chat_get_messages(
    id_login_current: Optional[str] = "",
    selected_workrooms_json: Union[str, list[str]] = "[]",
    latest_count: int = 20,
    options: int = 2,
    kind: int = 1,
    latest: bool = True,
    select_all: bool = False,
    latest_id_chat: int = 0,
) -> str:
    """
    LEE MENSAJES DE UNO O VARIOS CANALES EN SOLIDSET.
    Endpoint: POST /Chat/ChatMessages

    selected_workrooms_json acepta:
    - lista JSON en string: "[\"uuid\"]"
    - lista nativa: ["uuid"]
    """
    rooms: list[Any]
    if isinstance(selected_workrooms_json, list):
        rooms = selected_workrooms_json
    else:
        try:
            rooms = json.loads(selected_workrooms_json)
        except Exception as exc:
            return f"Error: selected_workrooms_json no es JSON válido: {exc}"

    if not isinstance(rooms, list) or not rooms:
        return (
            "Error: selected_workrooms_json debe contener al menos un ID de canal. "
            "Ejemplo válido: [\"debf64b2-3b3e-eb11-870c-d850e63f5833\"]"
        )

    room_ids = [str(item).strip() for item in rooms if str(item).strip()]
    if not room_ids:
        return "Error: selected_workrooms_json no contiene IDs de canal válidos."

    login_current = (id_login_current or "").strip()
    if login_current.lower() in {"current_user_id", "current-user-id", "user_id", ""}:
        login_current = ""

    params: dict[str, Any] = {
        "Options": int(options),
        "Kind": int(kind),
        "Latest": str(bool(latest)).lower(),
        "SelectAll": str(bool(select_all)).lower(),
        "LatestCount": int(latest_count),
        "LatestIDChat": int(latest_id_chat),
    }
    if login_current:
        params["IDLoginCurrent"] = login_current
    _solidset_add_indexed_params(params, "SelectedWorkRooms", room_ids)

    response, base, error = _solidset_request_authenticated(
        method="POST",
        endpoint="/Chat/ChatMessages",
        params=params,
    )
    if response is None:
        return f"Error consultando mensajes de chat: {error or 'sin detalle'}"
    return (
        f"status={response.status_code}; method=POST; "
        f"url={base}/Chat/ChatMessages; body={_solidset_preview_response(response)}"
    )


@tool
def solidset_chat_get_tasks_for_channel(
    id_workroom: str,
    running_states_json: Union[str, list[int]] = "[]",
    min_priority: int = 1,
    max_priority: int = 200,
    request_kind: int = 1,
) -> str:
    """
    CONSULTA TAREAS DE UN CANAL EN SOLIDSET.
    Endpoint: POST /Chat/GetTasksForChannelV2Form

    running_states_json acepta:
    - lista JSON en string: "[2723,2724,2725]"
    - lista nativa: [2723, 2724, 2725]
    """
    if isinstance(running_states_json, list):
        running_states = running_states_json
    else:
        try:
            running_states = json.loads(running_states_json)
        except Exception as exc:
            return f"Error: running_states_json no es JSON válido: {exc}"

    if not isinstance(running_states, list) or not running_states:
        return "Error: running_states_json debe contener al menos un estado."

    params: dict[str, Any] = {
        "MinPriority": int(min_priority),
        "MaxPriority": int(max_priority),
        "RequestKind": int(request_kind),
        "Assignments": 0,
        "Following": 0,
        "SelectAll": "false",
        "OrderBy[0][Field]": 4,
        "OrderBy[0][SortOrder]": 1,
        "MinImportance": 1,
        "MaxImportance": 55,
        "MinComplexity": 1,
        "MaxComplexity": 55,
        "Opts[0].FullVoteData": "true",
    }
    _solidset_add_indexed_params(params, "RunningStates", [str(int(v)) for v in running_states])
    _solidset_add_indexed_params(params, "IDWorkRooms", [id_workroom])

    response, base, error = _solidset_request_authenticated(
        method="POST",
        endpoint="/Chat/GetTasksForChannelV2Form",
        params=params,
    )
    if response is None:
        return f"Error consultando tareas del canal: {error or 'sin detalle'}"
    return (
        f"status={response.status_code}; method=POST; "
        f"url={base}/Chat/GetTasksForChannelV2Form; body={_solidset_preview_response(response)}"
    )


@tool
def solidset_point_get_task_info(id_task: str) -> str:
    """
    OBTIENE EL DETALLE DE UNA TAREA DE POINT.
    Endpoint: GET /Point/GetTaskInfo
    """
    params = {"idTask": (id_task or "").strip()}
    if not params["idTask"]:
        return "Error: id_task es obligatorio."

    response, base, error = _solidset_request_authenticated(
        method="GET",
        endpoint="/Point/GetTaskInfo",
        params=params,
    )
    if response is None:
        return f"Error consultando task info: {error or 'sin detalle'}"
    return (
        f"status={response.status_code}; method=GET; "
        f"url={base}/Point/GetTaskInfo; body={_solidset_preview_response(response)}"
    )


@tool
def solidset_point_get_activity_info(id_activity: int, id_module: str) -> str:
    """
    OBTIENE INFORMACIÓN DETALLADA DE UNA ACTIVIDAD DE POINT.
    Endpoint: GET /Point/GetActivityInfo
    """
    module_id = (id_module or "").strip()
    if not module_id:
        return "Error: id_module es obligatorio."

    params = {
        "idActivity": int(id_activity),
        "idModule": module_id,
    }
    response, base, error = _solidset_request_authenticated(
        method="GET",
        endpoint="/Point/GetActivityInfo",
        params=params,
    )
    if response is None:
        return f"Error consultando activity info: {error or 'sin detalle'}"
    return (
        f"status={response.status_code}; method=GET; "
        f"url={base}/Point/GetActivityInfo; body={_solidset_preview_response(response)}"
    )


@tool
def solidset_point_read_tasks(
    resource_id: str,
    read_activities: bool = True,
    only_tasks_assigned_to_me: bool = False,
    with_participants: bool = True,
    only_activities: bool = True,
    third_parties: bool = False,
    planned_start_date: Optional[str] = None,
    planned_end_date: Optional[str] = None,
) -> str:
    """
    LEE TAREAS/ACTIVIDADES DE POINT PARA UN RECURSO.
    Endpoint: GET /Point/ReadPointTask
    """
    rid = (resource_id or "").strip()
    if not rid:
        return "Error: resource_id es obligatorio."

    params: dict[str, Any] = {
        "readActivities": str(bool(read_activities)).lower(),
        "res[0]": rid,
        "onlyTasksAssignedToMe": str(bool(only_tasks_assigned_to_me)).lower(),
        "wPart": str(bool(with_participants)).lower(),
        "onlyAct": str(bool(only_activities)).lower(),
        "thirdParties": str(bool(third_parties)).lower(),
    }
    if planned_start_date:
        params["plannedStartDate"] = planned_start_date
    if planned_end_date:
        params["plannedEndDate"] = planned_end_date

    response, base, error = _solidset_request_authenticated(
        method="GET",
        endpoint="/Point/ReadPointTask",
        params=params,
    )
    if response is None:
        return f"Error leyendo tareas de Point: {error or 'sin detalle'}"
    return (
        f"status={response.status_code}; method=GET; "
        f"url={base}/Point/ReadPointTask; body={_solidset_preview_response(response)}"
    )


@tool
def solidset_vehicle_info(
    resource_id: str,
    with_last_logs: bool = True,
    last_log_page_size: int = 70,
    last_log_page: int = 1,
) -> str:
    """
    LEE DATOS DE VEHÍCULO Y ÚLTIMOS REGISTROS.
    Endpoint: GET /Vehicle/Info
    """
    rid = (resource_id or "").strip()
    if not rid:
        return "Error: resource_id es obligatorio."

    params = {
        "WithLastLogs": str(bool(with_last_logs)).lower(),
        "LastLogPageSize": int(last_log_page_size),
        "LastLogPage": int(last_log_page),
        "ResourceID": rid,
    }
    response, base, error = _solidset_request_authenticated(
        method="GET",
        endpoint="/Vehicle/Info",
        params=params,
    )
    if response is None:
        return f"Error consultando vehículo: {error or 'sin detalle'}"
    return (
        f"status={response.status_code}; method=GET; "
        f"url={base}/Vehicle/Info; body={_solidset_preview_response(response)}"
    )


@tool
def solidset_featureflag_get_resource_flags(resource_id: str) -> str:
    """
    LEE FEATURE FLAGS ACTIVAS PARA UN RECURSO.
    Endpoint: GET /FeatureFlag/GetResourceFeatureFlags
    """
    rid = (resource_id or "").strip()
    if not rid:
        return "Error: resource_id es obligatorio."

    response, base, error = _solidset_request_authenticated(
        method="GET",
        endpoint="/FeatureFlag/GetResourceFeatureFlags",
        params={"ResourceID": rid},
    )
    if response is None:
        return f"Error consultando feature flags del recurso: {error or 'sin detalle'}"
    return (
        f"status={response.status_code}; method=GET; "
        f"url={base}/FeatureFlag/GetResourceFeatureFlags; body={_solidset_preview_response(response)}"
    )


@tool
def solidset_featureflag_get_on() -> str:
    """
    LEE LAS FEATURE FLAGS ACTIVADAS GLOBALMENTE.
    Endpoint: GET /FeatureFlag/GetFeatureFlagsOn
    """
    response, base, error = _solidset_request_authenticated(
        method="GET",
        endpoint="/FeatureFlag/GetFeatureFlagsOn",
    )
    if response is None:
        return f"Error consultando feature flags activas: {error or 'sin detalle'}"
    return (
        f"status={response.status_code}; method=GET; "
        f"url={base}/FeatureFlag/GetFeatureFlagsOn; body={_solidset_preview_response(response)}"
    )


# ---------------------------------------------------------------------------
# 3. TOOLS: Telemetría y Diagnóstico CNC (MEJORADAS)
# ---------------------------------------------------------------------------
@tool
def get_cnc_telemetry() -> dict:
    """
    CONSULTA TELEMETRÍA EN TIEMPO REAL DE LA CNC HARTFORD.
    
    CUÁNDO USARLA:
    - El usuario pregunta "¿Cómo está la máquina?"
    - El usuario pregunta "¿Qué alarmas tiene?"
    - El usuario pregunta por RPM, temperatura o velocidad de avance
    
    DEVUELVE:
    - Estado de la máquina (OPERATIONAL/STOPPED/MAINTENANCE)
    - Velocidad del husillo (RPM)
    - Velocidad de avance (mm/min)
    - Potencia del husillo (%)
    - Alarmas activas
    """
    return {
        "status": "OPERATIONAL",
        "spindle_speed_rpm": 3200,
        "feed_rate_mm_min": 450,
        "spindle_power_pct": 88.5,
        "active_alarms": ["ALARM_102: Overload Spindle Warning"],
    }


@tool
def recommend_cnc_action(action: str, parameter: str, value: str) -> str:
    """
    REGISTRA UNA ACCIÓN CORRECTIVA PARA LA CNC.
    
    CUÁNDO USARLA:
    - Cuando el usuario pide cambiar un parámetro de la máquina
    - Cuando el usuario dice "Reduce la velocidad" o "Sube la temperatura"
    - Cuando se detecta una alarma y se necesita ajustar algo
    
    NOTA: Esta herramienta SOLO registra la acción, no la ejecuta físicamente.
    """
    return f"✅ Acción '{action}' con parámetro '{parameter}={value}' registrada en el sistema de mantenimiento."


@tool
def analyze_pcm_audio_diagnostic(file_path: str) -> str:
    """
    ANALIZA ARCHIVOS DE AUDIO .PCM DE LA CNC HARTFORD.
    
    CUÁNDO USARLA:
    - El usuario sube un archivo .pcm y pide diagnóstico
    - El usuario menciona "ruido extraño en la máquina"
    - El usuario pregunta por análisis acústico
    
    DEVUELVE:
    - Energía RMS (indica nivel de vibración)
    - Centroide espectral (frecuencia dominante)
    - Coincidencias con patrones conocidos (RAG)
    """
    try:
        features = extract_audio_features(file_path)
        rag_matches = get_rag_context(features["text_summary"])
        return (
            f"🔊 RESULTADOS DEL ANÁLISIS ACÚSTICO:\n\n"
            f"📁 Archivo: {features['file_name']}\n"
            f"📊 Energía RMS: {features['rms_energy']:.4f} (indica nivel de vibración)\n"
            f"🎵 Centroide Espectral: {features['spectral_centroid']:.2f} Hz\n"
            f"📚 Patrones coincidentes en la base de conocimiento:\n{rag_matches}"
        )
    except Exception as e:
        return f"Error al procesar el archivo de audio: {str(e)}"


@tool
def learn_new_fact(fact_description: str, category: str = "general") -> str:
    """
    GUARDA NUEVO CONOCIMIENTO EN LA BASE VECTORIAL.
    
    CUÁNDO USARLA:
    - El usuario dice "Aprende esto: ..."
    - El usuario comparte una observación importante sobre la máquina
    - El usuario enseña una regla de mantenimiento
    
    REGLAS:
    - SIEMPRE usa 'general' como categoría por defecto
    - Usa 'mantenimiento' si es sobre reparaciones
    - Usa 'operacion' si es sobre procedimientos de trabajo
    """
    try:
        client = QdrantClient(url=settings.VECTOR_DB_URL)
        embeddings = OllamaEmbeddings(
            base_url=settings.OLLAMA_BASE_URL,
            model=settings.EMBEDDING_MODEL_NAME,
        )

        ensure_vector_collection(client, settings.VECTOR_COLLECTION_NAME, embeddings)

        vector = embeddings.embed_query(fact_description)

        # 🚀 MEJORA: ID único basado en hash para evitar duplicados
        content_hash = hashlib.md5(fact_description.encode()).hexdigest()
        point_id = str(uuid.UUID(content_hash))

        point = PointStruct(
            id=point_id,
            vector=vector,
            payload={
                "page_content": fact_description,
                "category": category,
                "source": "operator_learning",
                "timestamp": str(uuid.uuid4()),
            },
        )

        client.upsert(collection_name=settings.VECTOR_COLLECTION_NAME, points=[point])
        return f"✅ Aprendizaje registrado correctamente: '{fact_description}'"
    except Exception as e:
        return f"Error al registrar el aprendizaje: {str(e)}"


# ---------------------------------------------------------------------------
# 4. TOOLS: Generación de documentos (Word / Excel / PDF)
# ---------------------------------------------------------------------------
@tool
def create_word_document(
    title: str,
    content: str,
    file_name: Optional[str] = None,
    document_kind: Optional[str] = None,
) -> str:
    """
    CREA UN DOCUMENTO WORD (.docx) con un título y contenido libre.

    CUÁNDO USARLA:
    - El usuario pida "hazme un Word" o "genera un informe en Word"
    - Se necesite exportar una respuesta a formato editable
    """
    try:
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.shared import Pt

        stem = _safe_file_stem(file_name or title, "documento")
        path = _generated_docs_dir() / f"{stem}_{_timestamp_suffix()}.docx"

        kind = (document_kind or _infer_document_kind(title, content)).strip() or "Documento"
        lines = _normalize_document_lines(content)

        doc = Document()
        header = doc.add_heading(title.strip() or "Documento", level=1)
        header.alignment = WD_ALIGN_PARAGRAPH.CENTER

        meta = doc.add_paragraph()
        meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
        meta_run = meta.add_run(f"Tipo: {kind} | Fecha: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
        meta_run.italic = True
        meta_run.font.size = Pt(10)

        doc.add_paragraph("")
        for line in lines:
            stripped = line.strip()
            if not stripped:
                doc.add_paragraph("")
                continue

            # Títulos de sección: "Sección:" o líneas totalmente en mayúsculas.
            is_section_title = stripped.endswith(":") or (stripped.upper() == stripped and len(stripped) <= 60)
            if is_section_title:
                p = doc.add_paragraph(stripped)
                p.runs[0].bold = True
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                continue

            # Viñetas básicas.
            if stripped.startswith(("- ", "• ", "* ")):
                p = doc.add_paragraph(stripped[2:].strip(), style="List Bullet")
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                continue

            p = doc.add_paragraph(stripped)
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        doc.save(path)

        return f"✅ Documento Word creado correctamente en: {path}"
    except Exception as e:
        return f"Error creando documento Word: {str(e)}"


@tool
def create_excel_document(
    title: str,
    rows_json: str,
    sheet_name: str = "Datos",
    file_name: Optional[str] = None,
    document_kind: Optional[str] = None,
) -> str:
    """
    CREA UN ARCHIVO EXCEL (.xlsx) a partir de filas en JSON.

    Formato esperado en rows_json:
    - Lista de objetos: [{"columna":"valor"}, {"columna":"valor2"}]
    """
    try:
        from openpyxl.styles import Alignment, Font, PatternFill, Border, Side
        from openpyxl.utils import get_column_letter

        stem = _safe_file_stem(file_name or title, "reporte")
        safe_sheet = _safe_file_stem(sheet_name, "Datos")[:31]
        path = _generated_docs_dir() / f"{stem}_{_timestamp_suffix()}.xlsx"
        kind = (document_kind or _infer_document_kind(title, rows_json)).strip() or "Documento"

        parsed = json.loads(rows_json) if rows_json else []
        if isinstance(parsed, dict):
            parsed = [parsed]
        if not isinstance(parsed, list):
            return "Error creando Excel: rows_json debe ser una lista JSON de filas u objeto JSON."

        df = pd.DataFrame(parsed)
        if df.empty:
            df = pd.DataFrame([{"mensaje": "Sin datos para exportar"}])

        with pd.ExcelWriter(path, engine="openpyxl") as writer:
            start_row = 3
            df.to_excel(writer, index=False, sheet_name=safe_sheet, startrow=start_row)
            worksheet = writer.sheets[safe_sheet]

            # Encabezado formal.
            worksheet.merge_cells(start_row=1, start_column=1, end_row=1, end_column=max(1, len(df.columns)))
            worksheet["A1"] = title.strip() or "Reporte"
            worksheet["A2"] = f"Tipo: {kind} | Fecha: {datetime.now().strftime('%Y-%m-%d %H:%M')}"
            worksheet["A1"].font = Font(bold=True, size=14)
            worksheet["A2"].font = Font(italic=True, size=10)

            header_row = start_row + 1
            thin = Side(style="thin", color="D9D9D9")
            for col_idx in range(1, len(df.columns) + 1):
                header_cell = worksheet.cell(row=header_row, column=col_idx)
                header_cell.font = Font(bold=True, color="FFFFFF")
                header_cell.fill = PatternFill("solid", fgColor="1F4E78")
                header_cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
                header_cell.border = Border(top=thin, left=thin, right=thin, bottom=thin)

            for row_idx in range(header_row + 1, header_row + 1 + len(df.index)):
                for col_idx in range(1, len(df.columns) + 1):
                    cell = worksheet.cell(row=row_idx, column=col_idx)
                    cell.alignment = Alignment(horizontal="left", vertical="top", wrap_text=True)
                    cell.border = Border(top=thin, left=thin, right=thin, bottom=thin)

            for idx, col in enumerate(df.columns, start=1):
                max_len = max(len(str(col)), *(len(str(v)) for v in df[col].astype(str).tolist()))
                letter = get_column_letter(idx)
                worksheet.column_dimensions[letter].width = min(max(12, max_len + 2), 50)

            worksheet.freeze_panes = worksheet.cell(row=header_row + 1, column=1)
            worksheet.auto_filter.ref = worksheet.dimensions

        return f"✅ Archivo Excel creado correctamente en: {path}"
    except json.JSONDecodeError:
        return "Error creando Excel: rows_json no tiene formato JSON válido."
    except Exception as e:
        return f"Error creando Excel: {str(e)}"


@tool
def create_pdf_document(
    title: str,
    content: str,
    file_name: Optional[str] = None,
    document_kind: Optional[str] = None,
) -> str:
    """
    CREA UN DOCUMENTO PDF (.pdf) con un título y contenido textual.

    CUÁNDO USARLA:
    - El usuario pida "hazme un PDF" o "exporta este resumen en PDF"
    """
    try:
        from reportlab.lib.pagesizes import letter
        from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
        from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY, TA_LEFT
        from reportlab.lib import colors
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, HRFlowable

        stem = _safe_file_stem(file_name or title, "documento")
        path = _generated_docs_dir() / f"{stem}_{_timestamp_suffix()}.pdf"
        kind = (document_kind or _infer_document_kind(title, content)).strip() or "Documento"
        lines = _normalize_document_lines(content)

        doc = SimpleDocTemplate(
            str(path),
            pagesize=letter,
            leftMargin=54,
            rightMargin=54,
            topMargin=54,
            bottomMargin=54,
        )

        styles = getSampleStyleSheet()
        title_style = ParagraphStyle(
            "DocTitle",
            parent=styles["Title"],
            alignment=TA_CENTER,
            fontName="Helvetica-Bold",
            fontSize=16,
            leading=20,
            spaceAfter=6,
        )
        meta_style = ParagraphStyle(
            "DocMeta",
            parent=styles["Normal"],
            alignment=TA_CENTER,
            fontName="Helvetica-Oblique",
            fontSize=9,
            textColor=colors.HexColor("#555555"),
            leading=12,
            spaceAfter=10,
        )
        section_style = ParagraphStyle(
            "DocSection",
            parent=styles["Heading3"],
            alignment=TA_LEFT,
            fontName="Helvetica-Bold",
            fontSize=11,
            leading=14,
            spaceBefore=8,
            spaceAfter=4,
        )
        body_style = ParagraphStyle(
            "DocBody",
            parent=styles["Normal"],
            alignment=TA_JUSTIFY,
            fontName="Helvetica",
            fontSize=10.5,
            leading=14,
            spaceAfter=6,
        )
        bullet_style = ParagraphStyle(
            "DocBullet",
            parent=body_style,
            leftIndent=14,
            bulletIndent=2,
            spaceAfter=4,
        )

        story = [
            Paragraph((title or "Documento").strip(), title_style),
            Paragraph(f"Tipo: {kind} | Fecha: {datetime.now().strftime('%Y-%m-%d %H:%M')}", meta_style),
            HRFlowable(width="100%", thickness=0.8, color=colors.HexColor("#D0D7DE"), spaceAfter=10),
        ]

        for raw in lines:
            stripped = raw.strip()
            if not stripped:
                story.append(Spacer(1, 4))
                continue

            safe = stripped.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
            is_section_title = safe.endswith(":") or (safe.upper() == safe and len(safe) <= 60)
            if is_section_title:
                story.append(Paragraph(safe, section_style))
                continue

            if safe.startswith(("- ", "• ", "* ")):
                story.append(Paragraph(safe[2:].strip(), bullet_style, bulletText="•"))
                continue

            story.append(Paragraph(safe, body_style))

        doc.build(story)
        return f"✅ Documento PDF creado correctamente en: {path}"
    except Exception as e:
        return f"Error creando PDF: {str(e)}"
